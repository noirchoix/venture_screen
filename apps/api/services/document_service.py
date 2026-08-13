from __future__ import annotations

import hashlib
import io
import re
import zipfile
from pathlib import PurePosixPath

from docx import Document
from pypdf import PdfReader

from core.config import settings


TEXT_EXTENSIONS = {
    '.txt', '.md', '.markdown', '.rst', '.csv', '.json', '.yaml', '.yml', '.toml',
    '.py', '.js', '.ts', '.tsx', '.jsx', '.svelte', '.java', '.go', '.rs', '.sql',
    '.html', '.css', '.scss', '.env.example',
}


class DocumentError(ValueError):
    pass


class DocumentService:
    def extract(self, filename: str, data: bytes, media_type: str = '') -> tuple[str, list[str], str]:
        if len(data) > settings.max_upload_bytes:
            raise DocumentError(f'File exceeds {settings.max_upload_bytes} byte upload limit.')
        sha = hashlib.sha256(data).hexdigest()
        lower = filename.lower()
        warnings: list[str] = []
        if lower.endswith('.pdf'):
            text = self._pdf(data, warnings)
        elif lower.endswith('.docx'):
            text = self._docx(data)
        elif lower.endswith('.zip'):
            text = self._zip(data, warnings)
        else:
            text = self._text(data, warnings)
        text = self._normalize(text)
        if len(text) > settings.max_document_chars:
            warnings.append(f'Content truncated to {settings.max_document_chars} characters for bounded screening context.')
            text = text[: settings.max_document_chars]
        if not text.strip():
            warnings.append('No machine-readable text was extracted from this file.')
        return text, warnings, sha

    def _text(self, data: bytes, warnings: list[str]) -> str:
        for encoding in ('utf-8', 'utf-8-sig', 'latin-1'):
            try:
                return data.decode(encoding)
            except UnicodeDecodeError:
                continue
        warnings.append('Text decoding used replacement characters.')
        return data.decode('utf-8', errors='replace')

    def _pdf(self, data: bytes, warnings: list[str]) -> str:
        reader = PdfReader(io.BytesIO(data))
        pages = []
        for index, page in enumerate(reader.pages):
            try:
                pages.append(page.extract_text() or '')
            except Exception as exc:  # parser resilience, not silent failure
                warnings.append(f'PDF page {index + 1} could not be parsed: {type(exc).__name__}.')
        if pages and sum(bool(x.strip()) for x in pages) < max(1, len(pages) // 3):
            warnings.append('PDF appears image-heavy; OCR is not performed by this service.')
        return '\n\n'.join(pages)

    def _docx(self, data: bytes) -> str:
        doc = Document(io.BytesIO(data))
        chunks = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                chunks.append(' | '.join(cell.text.strip() for cell in row.cells))
        return '\n'.join(chunks)

    def _zip(self, data: bytes, warnings: list[str]) -> str:
        parts: list[str] = []
        count = 0
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            for info in sorted(archive.infolist(), key=lambda x: x.filename):
                if info.is_dir() or info.file_size > 1_500_000:
                    continue
                path = PurePosixPath(info.filename)
                if path.is_absolute() or '..' in path.parts:
                    warnings.append(f'Skipped unsafe ZIP path: {info.filename}')
                    continue
                suffix = path.suffix.lower()
                if suffix not in TEXT_EXTENSIONS and path.name not in {'Dockerfile', 'Makefile', 'README'}:
                    continue
                if any(part in {'node_modules', '.git', '.svelte-kit', 'dist', 'build', '__pycache__'} for part in path.parts):
                    continue
                raw = archive.read(info)
                try:
                    text = raw.decode('utf-8')
                except UnicodeDecodeError:
                    continue
                parts.append(f'\n--- FILE: {info.filename} ---\n{text[:18000]}')
                count += 1
                if sum(len(x) for x in parts) >= settings.max_document_chars:
                    warnings.append('Repository ZIP context reached the configured character boundary.')
                    break
        warnings.append(f'Extracted {count} text/code files from repository ZIP.')
        return ''.join(parts)

    def _normalize(self, text: str) -> str:
        text = text.replace('\x00', '')
        text = re.sub(r'\r\n?', '\n', text)
        text = re.sub(r'\n{4,}', '\n\n\n', text)
        return text.strip()
